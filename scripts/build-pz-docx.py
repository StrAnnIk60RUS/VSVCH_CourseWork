# -*- coding: utf-8 -*-
"""
Сборка пояснительной записки VSVH в формате .docx.
Требует: pip install python-docx

Запуск из корня репозитория:
  python scripts/build-pz-docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DRAFT_PATH = ROOT / "docs" / "poyasnitelnaya-zapiska-draft.md"
OUTPUT_PATH = ROOT / "docs" / "poyasnitelnaya-zapiska.docx"
FALLBACK_OUTPUT_PATH = ROOT / "docs" / "poyasnitelnaya-zapiska-expanded.docx"

FONT = "Times New Roman"
INDENT_CM = 1.25

sys.path.insert(0, str(ROOT / "scripts"))
from pz_tables import ASSIGNMENT_TABLE, TABLE_1_1, TABLE_1_2, TABLE_1_3  # noqa: E402


def set_run_font(run, size_pt: int, bold: bool = False, italic: bool = False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_format(
    para,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    line_spacing: float = 1.15,
    first_line_indent_cm: float | None = INDENT_CM,
    space_after_pt: float = 0,
    space_before_pt: float = 0,
):
    pf = para.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = line_spacing_rule
    if line_spacing_rule == WD_LINE_SPACING.MULTIPLE:
        pf.line_spacing = line_spacing
    elif line_spacing_rule == WD_LINE_SPACING.SINGLE:
        pf.line_spacing = 1.0
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    else:
        pf.first_line_indent = Cm(0)
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(space_before_pt)


def add_text_block(doc, text: str, style: str):
    if style == "h1":
        para = doc.add_paragraph()
        set_paragraph_format(
            para,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_rule=WD_LINE_SPACING.SINGLE,
            first_line_indent_cm=None,
            space_after_pt=12,
        )
        run = para.add_run(text)
        set_run_font(run, 16, bold=True)
    elif style == "h2":
        para = doc.add_paragraph()
        set_paragraph_format(
            para,
            line_spacing_rule=WD_LINE_SPACING.SINGLE,
            first_line_indent_cm=INDENT_CM,
            space_before_pt=6,
            space_after_pt=6,
        )
        run = para.add_run(text)
        set_run_font(run, 16, bold=True)
    elif style == "h3":
        para = doc.add_paragraph()
        set_paragraph_format(para, first_line_indent_cm=INDENT_CM, space_before_pt=4, space_after_pt=4)
        run = para.add_run(text)
        set_run_font(run, 14, bold=True)
    elif style == "title":
        para = doc.add_paragraph()
        set_paragraph_format(
            para,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_rule=WD_LINE_SPACING.SINGLE,
            first_line_indent_cm=None,
            space_after_pt=0,
        )
        run = para.add_run(text)
        set_run_font(run, 14, bold=("МИНИСТЕРСТВО" in text or "КУРСОВАЯ" in text or "ВСВШ" in text))
    elif style == "fig_caption":
        para = doc.add_paragraph()
        set_paragraph_format(
            para,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_rule=WD_LINE_SPACING.SINGLE,
            first_line_indent_cm=None,
            space_before_pt=6,
            space_after_pt=6,
        )
        run = para.add_run(text)
        set_run_font(run, 12)
    elif style == "fig_placeholder":
        para = doc.add_paragraph()
        set_paragraph_format(
            para,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_rule=WD_LINE_SPACING.SINGLE,
            first_line_indent_cm=None,
        )
        run = para.add_run(text)
        set_run_font(run, 12, italic=True)
    else:
        for chunk in text.split("\n\n"):
            chunk = chunk.strip()
            if not chunk:
                continue
            para = doc.add_paragraph()
            set_paragraph_format(para)
            run = para.add_run(chunk)
            set_run_font(run, 12)


def add_table(doc, table_data: dict):
    cap = doc.add_paragraph()
    set_paragraph_format(
        cap,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
        first_line_indent_cm=INDENT_CM,
        space_after_pt=6,
    )
    run = cap.add_run(table_data["caption"])
    set_run_font(run, 12)

    headers = table_data["headers"]
    rows = table_data["rows"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                set_run_font(r, 11, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, 11)
    doc.add_paragraph()


def add_toc_field(doc):
    para = doc.add_paragraph()
    set_paragraph_format(
        para,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
        first_line_indent_cm=None,
        space_after_pt=12,
    )
    run = para.add_run("Содержание")
    set_run_font(run, 16, bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent_cm=None)
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    hint = doc.add_paragraph()
    set_paragraph_format(hint, first_line_indent_cm=None)
    hr = hint.add_run(
        "(После открытия в Word нажмите F9 или «Обновить поле», чтобы построить оглавление.)"
    )
    set_run_font(hr, 11, italic=True)
    doc.add_page_break()


def parse_draft(path: Path) -> list[tuple[str, str]]:
    if not path.exists():
        raise FileNotFoundError(f"Draft not found: {path}")
    text = path.read_text(encoding="utf-8")
    blocks: list[tuple[str, str]] = []
    tag_re = re.compile(r"^@(H1|H2|H3|P|FIG|TABLE|PAGE|TITLE|TOC)\s*(.*)$")
    current_style = None
    buffer: list[str] = []

    def flush():
        nonlocal buffer, current_style
        if current_style and buffer:
            blocks.append((current_style, "\n\n".join(buffer).strip()))
        buffer = []

    for line in text.splitlines():
        m = tag_re.match(line.strip())
        if m:
            flush()
            tag, rest = m.group(1), m.group(2).strip()
            if tag == "H1":
                current_style = "h1"
            elif tag == "H2":
                current_style = "h2"
            elif tag == "H3":
                current_style = "h3"
            elif tag == "P":
                current_style = "p"
            elif tag == "FIG":
                blocks.append(("fig", rest))
                current_style = None
            elif tag == "TABLE":
                blocks.append(("table", rest))
                current_style = None
            elif tag == "PAGE":
                blocks.append(("page", ""))
                current_style = None
            elif tag == "TITLE":
                current_style = "title"
                if rest:
                    buffer = [rest]
                    flush()
                    current_style = None
            elif tag == "TOC":
                blocks.append(("toc", ""))
                current_style = None
            else:
                current_style = "p"
            if rest and tag in ("H1", "H2", "H3"):
                buffer = [rest]
                flush()
                current_style = None
        elif line.strip() == "---":
            flush()
            current_style = None
        elif current_style:
            buffer.append(line)
        elif line.strip():
            buffer.append(line)
            current_style = "p"
    flush()
    return blocks


def build_title_pages(doc):
    lines = [
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "[НАЗВАНИЕ УНИВЕРСИТЕТА]",
        "[НАЗВАНИЕ ИНСТИТУТА / ФАКУЛЬТЕТА]",
        "",
        "Кафедра [название кафедры]",
        "",
        "",
        "КУРСОВАЯ РАБОТА",
        "по дисциплине «[Название дисциплины]»",
        "",
        "на тему:",
        "«Разработка web-приложения VSVH для изучения иностранных языков»",
        "",
        "",
        "",
        "Выполнил(а):",
        "студент(ка) группы [номер группы]",
        "[ФИО студента]",
        "",
        "Руководитель:",
        "[учёная степень, звание]",
        "[ФИО руководителя]",
        "",
        "",
        "[Город] – [2026]",
    ]
    for ln in lines:
        add_text_block(doc, ln, "title")
    doc.add_page_break()

    add_text_block(doc, "ЗАДАНИЕ", "h1")
    add_text_block(
        doc,
        "на курсовую работу студенту(ке) [ФИО] группы [номер] по теме "
        "«Разработка web-приложения VSVH для изучения иностранных языков».",
        "p",
    )
    add_table(doc, ASSIGNMENT_TABLE)
    add_text_block(
        doc,
        "Дата выдачи задания: «___» __________ 2026 г.\t\t"
        "Срок сдачи: «___» __________ 2026 г.",
        "p",
    )
    doc.add_page_break()


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    build_title_pages(doc)

    blocks = parse_draft(DRAFT_PATH)
    table_map = {
        "1.1": TABLE_1_1,
        "1.2": TABLE_1_2,
        "1.3": TABLE_1_3,
    }

    for style, content in blocks:
        if style == "toc":
            add_toc_field(doc)
        elif style == "page":
            doc.add_page_break()
        elif style == "fig":
            parts = content.split("|", 1)
            caption = parts[1].strip() if len(parts) > 1 else content
            add_text_block(doc, f"Рисунок {parts[0].strip()} – {caption}", "fig_caption")
            add_text_block(doc, f"[Место для рисунка {parts[0].strip()}]", "fig_placeholder")
        elif style == "table":
            key = content.strip()
            if key in table_map:
                add_table(doc, table_map[key])
        else:
            add_text_block(doc, content, style)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(OUTPUT_PATH))
        print(f"Saved: {OUTPUT_PATH}")
    except PermissionError:
        doc.save(str(FALLBACK_OUTPUT_PATH))
        print(f"Saved fallback (primary file is locked): {FALLBACK_OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
